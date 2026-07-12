from __future__ import annotations

import logging
from pathlib import Path
from utils.domain import Document

# Third-party binary parsing libraries
import pypdf
import docx

logger = logging.getLogger("SentinelIQ.Extractor")


class TextFileExtractor:
    """Extracts raw text content from standard flat plain-text documents."""
    
    def extract(self, path: Path) -> Document:
        try:
            content = path.read_text(encoding="utf-8")
            return Document(
                source_id=path.stem,
                title=path.name,
                content=content,
                metadata={"file_type": "text", "path": str(path)}
            )
        except Exception as e:
            raise IOError(f"Failed extracting raw text from target line: {e}")


class LogStreamExtractor:
    """Specialized parser reading operational logs line-by-line, filtering out noise."""
    
    def extract(self, path: Path) -> Document:
        valid_lines = []
        line_count = 0
        
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                line_count += 1
                clean_line = line.strip()
                # Production filter logic rules
                if not clean_line or "[DEBUG]" in clean_line:
                    continue
                if len(clean_line.split(" - ")) < 2:
                    logger.warning(f"Malformed operational stream record bypassed at line {line_count} inside {path.name}")
                    continue
                valid_lines.append(clean_line)
                
        return Document(
            source_id=f"LOG_{path.stem.upper()}",
            title=path.name,
            content="\n".join(valid_lines),
            metadata={"file_type": "log", "line_count": line_count, "path": str(path)}
        )


class PDFFileExtractor:
    """Production extractor decoding binary PDF structures page by page."""

    def extract(self, path: Path) -> Document:
        try:
            text_content = []
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            return Document(
                source_id=path.stem,
                title=path.name,
                content="\n".join(text_content),
                metadata={"file_type": "pdf", "pages": len(reader.pages), "path": str(path)}
            )
        except Exception as e:
            raise IOError(f"Failed extracting PDF binary text stream: {e}")


class DocxFileExtractor:
    """Production extractor parsing office OpenXML .docx compressed archives."""

    def extract(self, path: Path) -> Document:
        try:
            doc = docx.Document(path)
            # Pull text from all paragraphs within the asset archive
            paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
            
            return Document(
                source_id=path.stem,
                title=path.name,
                content="\n".join(paragraphs_text),
                metadata={"file_type": "docx", "paragraphs": len(paragraphs_text), "path": str(path)}
            )
        except Exception as e:
            raise IOError(f"Failed extracting word docx text matrix: {e}")